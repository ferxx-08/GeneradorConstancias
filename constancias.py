import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import tempfile
from docx2pdf import convert

st.set_page_config(page_title="Constancias UAEM", page_icon="📄")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RUTA_CSV = os.path.join(BASE_DIR, "nombres.csv")

def cargar_nombres():
    if not os.path.exists(RUTA_CSV):
        return []
    with open(RUTA_CSV, "r", encoding="utf-8") as f:
        lineas = f.readlines()
    return [l.strip() for l in lineas if l.strip() and l.strip() != "nombre"]

def guardar_nombre(nuevo):
    if not os.path.exists(RUTA_CSV):
        with open(RUTA_CSV, "w", encoding="utf-8") as f:
            f.write("nombre\n")
    nombres = cargar_nombres()
    if nuevo not in nombres:
        with open(RUTA_CSV, "a", encoding="utf-8") as f:
            f.write(nuevo + "\n")
        return True
    return False

def limpiar_campos():
    st.session_state["motivo"] = ""
    st.session_state["motivo_guardado"] = ""
    st.session_state["nombre_manual"] = ""
    if "fecha_manual" in st.session_state:
        st.session_state["fecha_manual"] = ""

if "input_key" not in st.session_state:
    st.session_state["input_key"] = "input_1"

if "mensaje_ok" not in st.session_state:
    st.session_state["mensaje_ok"] = False

with st.sidebar:
    st.image("uaem logo.png", width=200)
    st.markdown("### UAEM Valle de Chalco")
    st.write("Herramienta Administrativa")
    st.markdown("---")

st.markdown("""
<style>
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #e8f5e9, #fdf6e3);
}
[data-testid="stSidebar"] {
    background-color: #fdf6e3;
}
h1 {
    color: #004a2f !important;
    border-bottom: 3px solid #b3a369;
}
h3 {
    color: #b3a369 !important;
}
div.stButton > button {
    background-color: #004a2f !important;
    color: white !important;
    border: 2px solid #b3a369 !important;
    border-radius: 10px;
}
input, textarea {
    border: 2px solid #b3a369 !important;
    border-radius: 8px !important;
}
*:focus {
    outline: none !important;
    box-shadow: none !important;
}
</style>
""", unsafe_allow_html=True)

st.title("Generador de Constancias")
st.subheader("Área de Investigación y Estudios Avanzados")

tipo_constancia = st.radio(
    "Tipo de constancia",
    ["CON FIRMA DE DIRECTORA", "SIN FIRMA DE DIRECTORA"]
)

lista_nombres = cargar_nombres()

nombre_lista = st.selectbox(
    "Selecciona un nombre:",
    options=lista_nombres,
    index=None, 
    placeholder="Selecciona un nombre de la lista..."
)

nombre_manual = st.text_input("O escribe el nombre:", key="nombre_manual")

nombre = nombre_manual if nombre_manual else nombre_lista

nuevo = st.text_input("Agregar nuevo nombre", key=st.session_state["input_key"])

if st.button("Guardar nombre"):
    if nuevo.strip():
        guardado = guardar_nombre(nuevo.strip())
        if guardado:
            st.session_state["mensaje_ok"] = True
            st.session_state["input_key"] = f"input_{datetime.now().timestamp()}"
            st.rerun()
        else:
            st.warning("Ese nombre ya existe")
    else:
        st.warning("Escribe un nombre")

if st.session_state["mensaje_ok"]:
    st.success("Nombre guardado correctamente")
    st.session_state["mensaje_ok"] = False

bloquear = st.checkbox("Bloquear motivo")

if not bloquear:
    motivo = st.text_area("Motivo:", key="motivo")
    if motivo:
        st.session_state["motivo_guardado"] = motivo
else:
    motivo = st.session_state["motivo_guardado"]
    st.text_area("Motivo:", value=motivo, disabled=True)

modo_fecha = st.radio("Tipo de fecha", ["Automatica", "Manual"])

if modo_fecha == "Manual":
    fecha = st.text_input("Fecha:", key="fecha_manual")
else:
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    hoy = datetime.now()
    fecha = f"a {hoy.day} de {meses[hoy.month-1]} de {hoy.year}"

st.button("Limpiar todos los campos", on_click=limpiar_campos)

if st.button("Generar constancia"):
    if not nombre or not motivo.strip() or not fecha.strip():
        st.warning("Completa todos los campos")
    else:
        try:
            ruta = os.path.join(BASE_DIR,
                "CON FIRMA DE DIRECTORA.docx" if tipo_constancia == "CON FIRMA DE DIRECTORA"
                else "SIN FIRMA DE DIRECTORA.docx"
            )

            doc = Document(ruta)

            for p in doc.paragraphs:
                if "{{NOMBRE}}" in p.text:
                    p.clear()
                    run = p.add_run(nombre)
                    run.font.name = "Arial"
                    run.font.size = Pt(14)
                    run.bold = True

                if "{{MOTIVO}}" in p.text:
                    p.clear()
                    run = p.add_run(motivo)
                    run.font.name = "Arial"
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if "{{FECHA}}" in p.text:
                    p.text = fecha

            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "temp.docx")
                pdf_path = os.path.join(tmpdir, "temp.pdf")

                doc.save(docx_path)

                try:
                    convert(docx_path, pdf_path)
                    with open(pdf_path, "rb") as f:
                        st.download_button("Descargar PDF", f.read(), file_name=f"Constancia_{nombre}.pdf")
                except:
                    with open(docx_path, "rb") as f:
                        st.download_button("Descargar Word", f.read(), file_name=f"Constancia_{nombre}.docx")

        except Exception as e:
            st.error(f"Error: {e}")